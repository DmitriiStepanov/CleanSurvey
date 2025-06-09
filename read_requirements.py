from docx import Document
import os

def read_docx_to_text(input_path):
    """
    Читает .docx файл и возвращает его содержимое в виде текста.
    """
    doc = Document(input_path)
    text = []
    
    # Читаем все параграфы
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
    
    # Читаем все таблицы
    for table in doc.tables:
        text.append("\n=== ТАБЛИЦА ===")
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            text.append(" | ".join(row_text))
        text.append("=== КОНЕЦ ТАБЛИЦЫ ===\n")
    
    return "\n".join(text)

def save_requirements():
    """
    Читает файл с требованиями и сохраняет его в текстовый файл.
    """
    input_file = "Задача 2 - рынок маркет ресерча.docx"
    output_file = "requirements.txt"
    
    if not os.path.exists(input_file):
        print(f"Ошибка: файл {input_file} не найден")
        return
    
    try:
        text = read_docx_to_text(input_file)
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"Требования успешно сохранены в {output_file}")
    except Exception as e:
        print(f"Ошибка при обработке файла: {e}")

if __name__ == "__main__":
    save_requirements() 