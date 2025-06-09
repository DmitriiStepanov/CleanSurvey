"""
doc_generator.py
Формирование нового clean_survey_MVP.docx на основе структуры анкеты.
"""

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from utils import set_run_color

def set_cell_border(cell, **kwargs):
    """
    Установка границ для ячейки таблицы
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ['top', 'left', 'bottom', 'right']:
        tag = f'{edge}'
        element = OxmlElement(f'w:{tag}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')  # 0.5pt
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'BFBFBF')
        
        tcBorders = tcPr.get_or_add_tcBorders()
        tcBorders.append(element)

def set_paragraph_background(paragraph, color):
    """
    Установка цвета фона для параграфа
    """
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    paragraph._p.get_or_add_pPr().append(shd)

def generate_clean_doc(survey_structure, output_path):
    """
    Принимает survey_structure (словарь из group_paragraphs) и создаёт .docx файл.
    """
    doc = Document()
    
    # Установка шрифта по умолчанию
    style = doc.styles['Normal']
    style.font.name = 'Montserrat'
    style.font.size = Pt(9)
    style.font.color.rgb = RGBColor(0, 0, 0)

    # 1. Общие инструкции
    general = survey_structure.get("general_instructions", [])
    if general:
        for line in general:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line.upper())
            set_run_color(run, (52, 116, 168))  # #3474a8

    # 2. Блоки анкеты
    for blk in survey_structure.get("blocks", []):
        # 2.1. Заголовок блока
        p_blk = doc.add_paragraph()
        p_blk.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_blk = p_blk.add_run(blk["header"])
        run_blk.font.bold = True
        run_blk.font.color.rgb = RGBColor(255, 255, 255)  # белый
        # Установка фона для параграфа
        set_paragraph_background(p_blk, "3474A8")

        # 2.2. Инструкции к блоку
        for instr in blk.get("instructions", []):
            p_instr = doc.add_paragraph()
            p_instr.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_instr = p_instr.add_run(instr.upper())
            set_run_color(run_instr, (52, 116, 168))

        # 2.3. Вопросы в блоке
        for q in blk.get("questions", []):
            # пустая строка для разделения
            doc.add_paragraph()

            # 2.3.1. Вопрос
            p_q = doc.add_paragraph()
            # Разделяем код вопроса и текст
            parts = q["text"].split(".", 1)
            if len(parts) > 1:
                code = parts[0] + "."
                text = parts[1].strip()
                run_code = p_q.add_run(code)
                run_code.font.bold = True
                p_q.add_run(text)
            else:
                p_q.add_run(q["text"])

            # 2.3.2. Инструкции к вопросу
            for qi in q.get("instructions", []):
                p_qi = doc.add_paragraph()
                p_qi.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run_qi = p_qi.add_run(qi.upper())
                set_run_color(run_qi, (52, 116, 168))

            # 2.3.3. Варианты ответов
            opts = q.get("answer_options", [])
            if opts:
                doc.add_paragraph()  # отступ
                p_opts = doc.add_paragraph("ВАРИАНТЫ ОТВЕТА:")
                run_opts = p_opts.runs[0]
                run_opts.font.bold = True
                set_run_color(run_opts, (52, 116, 168))
                
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = "КОД"
                hdr[1].text = "ФОРМУЛИРОВКА"
                hdr[2].text = "ИНСТРУКЦИЯ"
                
                # Форматирование шапки
                for cell in hdr:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(52, 116, 168)
                    cell._tc.get_or_add_tcPr().get_or_add_shd().fill = "D9D9D9"
                    set_cell_border(cell)
                
                for opt in opts:
                    row_cells = table.add_row().cells
                    row_cells[0].text = opt["code"]
                    row_cells[1].text = opt["text"]
                    row_cells[2].text = opt["instruction"].upper()
                    for cell in row_cells:
                        set_cell_border(cell)

            # 2.3.4. Строки (матрица)
            rows = q.get("rows", [])
            if rows:
                doc.add_paragraph()  # отступ
                p_rows = doc.add_paragraph("СТРОКИ:")
                run_rows = p_rows.runs[0]
                run_rows.font.bold = True
                set_run_color(run_rows, (52, 116, 168))
                
                table_r = doc.add_table(rows=1, cols=3)
                table_r.style = 'Table Grid'
                hdr_r = table_r.rows[0].cells
                hdr_r[0].text = "КОД"
                hdr_r[1].text = "ФОРМУЛИРОВКА"
                hdr_r[2].text = "ИНСТРУКЦИЯ"
                
                # Форматирование шапки
                for cell in hdr_r:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(52, 116, 168)
                    cell._tc.get_or_add_tcPr().get_or_add_shd().fill = "D9D9D9"
                    set_cell_border(cell)
                
                for row in rows:
                    rc = table_r.add_row().cells
                    rc[0].text = row["code"]
                    rc[1].text = row["text"]
                    rc[2].text = row["instruction"].upper()
                    for cell in rc:
                        set_cell_border(cell)

    # 4. Сохранение
    doc.save(output_path) 