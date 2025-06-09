from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_document():
    doc = Document()
    
    # Общие инструкции
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Пожалуйста, заполните все поля анкеты. Отвечайте честно и полно.")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(52, 116, 168)  # #3474a8
    
    # Заголовок блока
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ДЕМОГРАФИЧЕСКИЕ ДАННЫЕ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(255, 255, 255)  # Белый
    p.paragraph_format.space_after = Pt(12)
    
    # Инструкция к блоку
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("В этом блоке укажите основную информацию о себе.")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(52, 116, 168)  # #3474a8
    
    # Вопрос
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Q1. Укажите ваш пол:")
    run.font.size = Pt(9)
    run.font.bold = True
    
    # Варианты ответов
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("ВАРИАНТЫ ОТВЕТА:")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(52, 116, 168)  # #3474a8
    
    # Создаем таблицу с правильным количеством строк
    data = [
        ('1', 'Мужской', 'Выберите один вариант'),
        ('2', 'Женский', 'Выберите один вариант'),
        ('3', 'Предпочитаю не указывать', 'Выберите один вариант')
    ]
    
    table = doc.add_table(rows=len(data) + 1, cols=3)  # +1 для заголовка
    table.style = 'Table Grid'
    
    # Заголовки таблицы
    headers = ['КОД', 'ФОРМУЛИРОВКА', 'ИНСТРУКЦИЯ']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(52, 116, 168)
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Данные
    for i, row in enumerate(data):
        for j, text in enumerate(row):
            cell = table.cell(i + 1, j)  # +1 потому что первая строка - заголовок
            cell.text = text
            if j == 2:  # Инструкция
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(52, 116, 168)
    
    # Сохраняем документ
    doc.save('examples/dirty_survey_simple.docx')
    print("Тестовый документ создан: examples/dirty_survey_simple.docx")

if __name__ == "__main__":
    create_test_document() 